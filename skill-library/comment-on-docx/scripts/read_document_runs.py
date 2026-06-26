"""
Helper script to read a Word document as numbered runs.
Displays all runs with their indices, making it easy to reference specific text for commenting.
Also extracts images from the document for visual review.
"""
from docx import Document
from lxml import etree
import sys
import os
import tempfile
import zipfile
from typing import Optional
from docx.text.paragraph import Paragraph as ParagraphCls
from docx.table import Table as TableCls

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
M = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'


def iter_all_runs(para, hyperlink_rels=None):
    """
    Yield (run_element, is_hyperlink, hyperlink_url) for every <w:r> in the paragraph,
    in document order, including runs nested inside <w:hyperlink> and <w:ins>.
    Skips runs inside <w:del> (proposed deletions / track changes).
    This shows the "all suggestions accepted" version of the document.

    If hyperlink_rels is provided (dict mapping r:id -> URL), hyperlink_url
    will be the resolved URL for runs inside <w:hyperlink> elements.
    """
    def _yield_runs(container):
        for child in container:
            tag = child.tag.split('}')[-1]
            if tag == 'r':
                yield child, False, None
            elif tag == 'hyperlink':
                rid = child.get(f'{R_NS}id')
                url = hyperlink_rels.get(rid) if hyperlink_rels and rid else None
                for inner in child.findall(f'{W}r'):
                    yield inner, True, url
            elif tag == 'ins':
                # Proposed insertions — recurse to pick up runs and hyperlinks
                yield from _yield_runs(child)
            # 'del' is implicitly skipped (proposed deletions)

    yield from _yield_runs(para._element)


def _is_toc_sdt(sdt_element):
    """Check if a structured document tag is a Table of Contents."""
    sdt_pr = sdt_element.find(f'{W}sdtPr')
    if sdt_pr is not None:
        if sdt_pr.find(f'{W}docPartObj') is not None:
            return True
    return False


def _iter_sdt_content(sdt_content, parent):
    """Yield (Paragraph, table_info) for paragraphs and tables inside SDT content."""
    for item in sdt_content:
        tag = item.tag.split('}')[-1]
        if tag == 'p':
            yield ParagraphCls(item, parent), None
        elif tag == 'tbl':
            tbl = TableCls(item, parent)
            num_rows = len(tbl.rows)
            num_cols = len(tbl.columns)
            for row_idx, row in enumerate(tbl.rows):
                seen_tc = set()
                for col_idx, cell in enumerate(row.cells):
                    tc_id = id(cell._tc)
                    if tc_id in seen_tc:
                        continue
                    seen_tc.add(tc_id)
                    for para in cell.paragraphs:
                        yield para, {
                            'row': row_idx,
                            'col': col_idx,
                            'num_rows': num_rows,
                            'num_cols': num_cols,
                        }


def _iter_document_paragraphs(doc):
    """
    Yield (Paragraph, table_info) for every paragraph in the document body,
    in document order, including paragraphs inside table cells and SDTs.

    table_info is None for body paragraphs, or a dict with row/col/dimensions
    for table cell paragraphs. Skips Table of Contents SDTs.
    """
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            yield ParagraphCls(child, body), None
        elif tag == 'tbl':
            tbl = TableCls(child, body)
            num_rows = len(tbl.rows)
            num_cols = len(tbl.columns)
            for row_idx, row in enumerate(tbl.rows):
                seen_tc = set()
                for col_idx, cell in enumerate(row.cells):
                    # Skip duplicate cells from merged cells
                    tc_id = id(cell._tc)
                    if tc_id in seen_tc:
                        continue
                    seen_tc.add(tc_id)
                    for para in cell.paragraphs:
                        yield para, {
                            'row': row_idx,
                            'col': col_idx,
                            'num_rows': num_rows,
                            'num_cols': num_cols,
                        }
        elif tag == 'sdt':
            # Skip Table of Contents SDTs (duplicate heading text)
            if _is_toc_sdt(child):
                continue
            sdt_content = child.find(f'{W}sdtContent')
            if sdt_content is not None:
                yield from _iter_sdt_content(sdt_content, body)


def extract_images(docx_path: str, output_dir: str = None) -> tuple:
    """
    Extract all images from a docx file and build relationship mapping.

    Returns:
        tuple: (output_dir, rel_id_to_filename dict)
    """
    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix='docx_images_')
    os.makedirs(output_dir, exist_ok=True)

    rel_id_to_filename = {}

    with zipfile.ZipFile(docx_path) as z:
        # Extract image files to flat directory
        for name in z.namelist():
            if name.startswith('word/media/'):
                filename = os.path.basename(name)
                with z.open(name) as src:
                    with open(os.path.join(output_dir, filename), 'wb') as dst:
                        dst.write(src.read())

        # Parse relationships to map rId -> filename
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path in z.namelist():
            rels_xml = z.read(rels_path)
            rels_tree = etree.fromstring(rels_xml)
            for rel in rels_tree:
                target = rel.get('Target', '')
                if target.startswith('media/'):
                    rel_id = rel.get('Id')
                    rel_id_to_filename[rel_id] = os.path.basename(target)

    return output_dir, rel_id_to_filename


def get_image_in_element(element, rel_id_to_filename: dict) -> Optional[str]:
    """
    Check if an XML element contains an embedded image and return its filename.
    Searches for <a:blip r:embed="rIdX"> anywhere in the element's descendants.
    """
    blips = element.findall(f'.//{A}blip')
    for blip in blips:
        embed = blip.get(f'{R_NS}embed')
        if embed and embed in rel_id_to_filename:
            return rel_id_to_filename[embed]
    return None


def get_paragraph_level_images(para_element, rel_id_to_filename: dict) -> list:
    """
    Find images in a paragraph that are NOT inside <w:r> or <w:hyperlink> elements.
    These are typically in <mc:AlternateContent> or other paragraph-level elements.
    """
    images = []
    for child in para_element:
        tag = child.tag.split('}')[-1]
        if tag in ('r', 'hyperlink'):
            continue
        filename = get_image_in_element(child, rel_id_to_filename)
        if filename:
            images.append(filename)
    return images


def parse_footnotes(docx_path: str) -> dict:
    """
    Parse footnotes from word/footnotes.xml.

    Returns:
        dict mapping footnote ID (str) to footnote text.
        Skips separator/continuation footnotes (IDs 0 and -1).
    """
    footnotes = {}
    with zipfile.ZipFile(docx_path) as z:
        if 'word/footnotes.xml' not in z.namelist():
            return footnotes
        fn_xml = z.read('word/footnotes.xml')
        fn_tree = etree.fromstring(fn_xml)
        for fn in fn_tree.findall(f'{W}footnote'):
            fn_id = fn.get(f'{W}id')
            # Skip separator footnotes (type="separator" or "continuationSeparator")
            fn_type = fn.get(f'{W}type')
            if fn_type in ('separator', 'continuationSeparator'):
                continue
            # Extract all text, handling <w:ins>/<w:del> the same way as the body
            text = ''.join(t.text or '' for t in fn.iter(f'{W}t'))
            if text.strip():
                footnotes[fn_id] = text.strip()
    return footnotes


def parse_endnotes(docx_path: str) -> dict:
    """
    Parse endnotes from word/endnotes.xml.

    Returns:
        dict mapping endnote ID (str) to endnote text.
        Skips separator/continuation endnotes.
    """
    endnotes = {}
    with zipfile.ZipFile(docx_path) as z:
        if 'word/endnotes.xml' not in z.namelist():
            return endnotes
        en_xml = z.read('word/endnotes.xml')
        en_tree = etree.fromstring(en_xml)
        for en in en_tree.findall(f'{W}endnote'):
            en_id = en.get(f'{W}id')
            en_type = en.get(f'{W}type')
            if en_type in ('separator', 'continuationSeparator'):
                continue
            text = ''.join(t.text or '' for t in en.iter(f'{W}t'))
            if text.strip():
                endnotes[en_id] = text.strip()
    return endnotes


def get_equations_in_paragraph(para_element) -> list:
    """
    Find <m:oMath> elements in a paragraph and return their text.
    Equations use <m:t> for text instead of <w:t>.
    """
    equations = []
    for child in para_element:
        tag = child.tag.split('}')[-1]
        if tag == 'oMath':
            eq_text = ''.join(t.text or '' for t in child.iter(f'{M}t'))
            if eq_text.strip():
                equations.append(eq_text.strip())
    return equations


def read_document_runs(docx_path: str) -> dict:
    """
    Read document and return all runs numbered for easy reference.
    Also extracts images and footnotes from the document.

    Returns:
        dict with keys:
            - 'runs': list of dicts with run info (para_idx, run_idx, text, bold, italic, is_hyperlink, image, footnote_id, table_info)
            - 'comments': list of existing comments with their anchored runs
            - 'total_runs': total number of runs
            - 'total_chars': total character count
            - 'total_tables': number of tables in the document
            - 'images': list of image info dicts (filename, path, para_idx, in_run)
            - 'image_dir': path to directory containing extracted images (or None)
            - 'footnotes': dict mapping footnote ID to text
    """
    doc = Document(docx_path)

    # Extract images if the document has any
    image_dir = None
    rel_id_to_filename = {}

    # Parse relationships from the docx zip
    hyperlink_rels = {}
    with zipfile.ZipFile(docx_path) as z:
        has_images = any(f.startswith('word/media/') for f in z.namelist())

        # Parse hyperlink relationships
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path in z.namelist():
            rels_xml = z.read(rels_path)
            rels_tree = etree.fromstring(rels_xml)
            for rel in rels_tree:
                if 'hyperlink' in rel.get('Type', '').lower():
                    hyperlink_rels[rel.get('Id')] = rel.get('Target', '')

    if has_images:
        image_dir, rel_id_to_filename = extract_images(docx_path)

    # Parse footnotes and endnotes
    footnotes = parse_footnotes(docx_path)
    endnotes = parse_endnotes(docx_path)

    # Count tables
    table_count = len(doc.element.body.findall(f'{W}tbl'))

    all_runs = []
    all_images = []
    all_equations = []  # list of {'text': str, 'para_idx': int}
    run_counter = 0
    total_chars = 0

    # Collect all runs from all paragraphs, including table cells and hyperlink runs
    for para_idx, (para, table_info) in enumerate(_iter_document_paragraphs(doc)):
        for run_elem, is_hyperlink, hyperlink_url in iter_all_runs(para, hyperlink_rels):
            rPr = run_elem.find(f'{W}rPr')
            text = run_elem.findtext(f'{W}t', default='')
            bold = rPr is not None and rPr.find(f'{W}b') is not None if rPr is not None else False
            italic = rPr is not None and rPr.find(f'{W}i') is not None if rPr is not None else False

            # Check for image in this run
            image_filename = get_image_in_element(run_elem, rel_id_to_filename) if rel_id_to_filename else None

            # Check for footnote/endnote references in this run
            fn_ref = run_elem.find(f'{W}footnoteReference')
            footnote_id = fn_ref.get(f'{W}id') if fn_ref is not None else None
            en_ref = run_elem.find(f'{W}endnoteReference')
            endnote_id = en_ref.get(f'{W}id') if en_ref is not None else None

            run_info = {
                'global_run_id': run_counter,
                'para_idx': para_idx,
                'text': text,
                'bold': bold,
                'italic': italic,
                'is_hyperlink': is_hyperlink,
                'hyperlink_url': hyperlink_url,
                'image': image_filename,
                'footnote_id': footnote_id,
                'endnote_id': endnote_id,
                'table_info': table_info,
            }
            all_runs.append(run_info)
            total_chars += len(text)

            if image_filename:
                all_images.append({
                    'filename': image_filename,
                    'path': os.path.join(image_dir, image_filename),
                    'para_idx': para_idx,
                    'in_run': run_counter,
                })

            run_counter += 1

        # Check for paragraph-level images not inside runs
        if rel_id_to_filename:
            para_images = get_paragraph_level_images(para._element, rel_id_to_filename)
            for img_filename in para_images:
                all_images.append({
                    'filename': img_filename,
                    'path': os.path.join(image_dir, img_filename),
                    'para_idx': para_idx,
                    'in_run': None,
                })

        # Check for equations (m:oMath) in the paragraph
        para_equations = get_equations_in_paragraph(para._element)
        for eq_text in para_equations:
            all_equations.append({
                'text': eq_text,
                'para_idx': para_idx,
            })

    # Read existing comments and find which paragraphs they're in
    existing_comments = []
    if hasattr(doc, 'comments') and doc.comments is not None:
        try:
            comment_locations = {}  # comment_id -> para_idx

            for para_idx, (para, _) in enumerate(_iter_document_paragraphs(doc)):
                para_elem = para._element
                # Look for commentRangeStart or commentReference in this paragraph
                for elem in para_elem.iter():
                    if elem.tag == f'{W}commentRangeStart' or elem.tag == f'{W}commentReference':
                        comment_id = elem.get(f'{W}id')
                        if comment_id is not None:
                            comment_locations[comment_id] = para_idx

            for i, comment in enumerate(doc.comments):
                comment_id_str = str(getattr(comment, 'comment_id', getattr(comment, 'id', i)))
                para_idx = comment_locations.get(comment_id_str)

                # python-docx .text misses runs inside <w:ins> (Google Docs exports).
                # Fall back to extracting all <w:t> text from the XML element directly.
                text = getattr(comment, 'text', '') or ''
                if not text.strip():
                    text = ''.join(
                        t.text or '' for t in comment._element.iter(f'{W}t')
                    )

                comment_info = {
                    'id': getattr(comment, 'comment_id', getattr(comment, 'id', i)),
                    'author': getattr(comment, 'author', 'Unknown'),
                    'text': text.strip(),
                    'para_idx': para_idx,
                }
                existing_comments.append(comment_info)
        except Exception as e:
            print(f"⚠️  Warning: Could not fully read all existing comments: {e}")

    return {
        'runs': all_runs,
        'comments': existing_comments,
        'total_runs': len(all_runs),
        'total_chars': total_chars,
        'total_tables': table_count,
        'images': all_images,
        'image_dir': image_dir,
        'footnotes': footnotes,
        'endnotes': endnotes,
        'equations': all_equations,
    }


def display_document_runs(docx_path: str) -> None:
    """Display document runs in a format easy for Claude to read."""

    print("=" * 80)
    print(f"READING: {docx_path}")
    print("=" * 80)

    result = read_document_runs(docx_path)

    print(f"\n📊 DOCUMENT STATISTICS:")
    print(f"   Total runs: {result['total_runs']}")
    print(f"   Total characters: {result['total_chars']:,}")
    print(f"   Existing comments: {len(result['comments'])}")
    print(f"   Images: {len(result['images'])}")
    print(f"   Footnotes: {len(result['footnotes'])}")
    print(f"   Endnotes: {len(result['endnotes'])}")
    print(f"   Equations: {len(result['equations'])}")
    print(f"   Tables: {result['total_tables']}")

    if result['footnotes']:
        print(f"\n📝 FOOTNOTES:")
        for fn_id, fn_text in sorted(result['footnotes'].items(), key=lambda x: int(x[0])):
            print(f"   [Footnote {fn_id}] {fn_text}")

    if result['endnotes']:
        print(f"\n📝 ENDNOTES:")
        for en_id, en_text in sorted(result['endnotes'].items(), key=lambda x: int(x[0])):
            print(f"   [Endnote {en_id}] {en_text}")

    if result['image_dir'] and result['images']:
        # Deduplicate by filename (same image can appear multiple times)
        seen = set()
        unique_images = []
        for img in result['images']:
            if img['filename'] not in seen:
                seen.add(img['filename'])
                unique_images.append(img)

        print(f"\n🖼️  EXTRACTED IMAGES (saved to {result['image_dir']}/):")
        for img in unique_images:
            para_info = f" (Paragraph {img['para_idx']}"
            if img['in_run'] is not None:
                para_info += f", Run {img['in_run']}"
            para_info += ")"
            print(f"   {img['path']}{para_info}")
        print(f"\n   ➡️  Use the Read tool to view each image file listed above for full document context.")

    if result['comments']:
        print(f"\n💬 EXISTING COMMENTS:")
        for comment in result['comments']:
            author = comment['author']
            text = comment['text']
            para_info = f" (Paragraph {comment['para_idx']})" if comment.get('para_idx') is not None else ""
            print(f"   [{author}]{para_info} {text}")

    print(f"\n📖 ALL RUNS (numbered for easy reference):")
    print("=" * 80)

    # Build set of paragraph-level images (not in any run) keyed by para_idx
    para_level_images = {}
    for img in result.get('images', []):
        if img['in_run'] is None:
            para_level_images.setdefault(img['para_idx'], []).append(img)

    # Build set of equations keyed by para_idx
    para_equations = {}
    for eq in result.get('equations', []):
        para_equations.setdefault(eq['para_idx'], []).append(eq)

    current_para = -1
    in_table = False
    for run_info in result['runs']:
        table_info = run_info.get('table_info')

        # Print paragraph separator when we move to a new paragraph
        if run_info['para_idx'] != current_para:
            # Show paragraph-level images from the previous paragraph
            if current_para in para_level_images:
                for img in para_level_images[current_para]:
                    print(f"         [IMAGE: {img['filename']}]")

            # Show equations from the previous paragraph
            if current_para in para_equations:
                for eq in para_equations[current_para]:
                    print(f"         [EQUATION: {eq['text']}]")

            # Handle table transitions
            if table_info and not in_table:
                in_table = True
                print(f"\n--- Table ({table_info['num_rows']} rows × {table_info['num_cols']} cols) ---")
            elif not table_info and in_table:
                in_table = False
                print(f"--- End Table ---")

            current_para = run_info['para_idx']
            if table_info:
                print(f"\n  --- Paragraph {current_para} [Row {table_info['row']}, Col {table_info['col']}] ---")
            else:
                print(f"\n--- Paragraph {current_para} ---")

        # Format the run display
        run_id = run_info['global_run_id']
        text = run_info['text']
        image = run_info.get('image')

        # Show formatting indicators
        formatting = []
        if run_info.get('is_hyperlink'):
            url = run_info.get('hyperlink_url')
            formatting.append(f'LINK: {url}' if url else 'LINK')
        if run_info['bold']:
            formatting.append('BOLD')
        if run_info['italic']:
            formatting.append('ITALIC')
        format_str = f" [{', '.join(formatting)}]" if formatting else ""

        # Check for footnote/endnote references
        footnote_id = run_info.get('footnote_id')
        endnote_id = run_info.get('endnote_id')

        # Display the run
        if not text and image:
            print(f"[Run {run_id}] [IMAGE: {image}]{format_str}")
        elif not text and footnote_id:
            print(f"[Run {run_id}] [FOOTNOTE {footnote_id}]")
        elif not text and endnote_id:
            print(f"[Run {run_id}] [ENDNOTE {endnote_id}]")
        elif not text:
            print(f"[Run {run_id}] [EMPTY]{format_str}")
        else:
            img_str = f" [IMAGE: {image}]" if image else ""
            fn_str = f" [FOOTNOTE {footnote_id}]" if footnote_id else ""
            en_str = f" [ENDNOTE {endnote_id}]" if endnote_id else ""
            display_text = text
            print(f"[Run {run_id}] {display_text}{format_str}{img_str}{fn_str}{en_str}")

    # Show paragraph-level images and equations for the last paragraph
    if current_para in para_level_images:
        for img in para_level_images[current_para]:
            print(f"         [IMAGE: {img['filename']}]")
    if current_para in para_equations:
        for eq in para_equations[current_para]:
            print(f"         [EQUATION: {eq['text']}]")

    # Close trailing table if document ends inside one
    if in_table:
        print(f"--- End Table ---")

    print("=" * 80)
    print(f"\n✅ Document read complete. Total runs: {result['total_runs']}")
    print(f"   (Make sure you see all runs from [Run 0] to [Run {result['total_runs'] - 1}])")

    if result['images']:
        print(f"\n🖼️  IMPORTANT: This document contains {len(result['images'])} image(s).")
        print(f"   Read the images from {result['image_dir']}/ to understand figures and diagrams.")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python read_document_runs.py <path_to_docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    display_document_runs(docx_path)
