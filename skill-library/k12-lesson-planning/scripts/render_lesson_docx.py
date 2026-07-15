#!/usr/bin/env python3
# Copyright 2026 Anthropic, PBC
# Copyright 2026 Learning Commons
# SPDX-License-Identifier: Apache-2.0

"""Render lesson JSON -> an editable .docx (the teacher-editable deliverable).

Same input schema and block vocabulary as render_lesson_html.py; consumes the same
expand_document() / theme / alias / callout-kind helpers from lesson_common, so the two
formats stay in sync by construction. Styling follows the same visual design principles:
minimal color, icon-prefixed callouts (no fill), B+W-safe.

Usage:
    python render_lesson_docx.py lesson.json -o lesson_plan.docx
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from lesson_common import (  # noqa: E402
    Theme, CALLOUT_KINDS, FILL_IN_CHARS, btype as _btype,
    resolve_callout_kind, answer_profile, build_header, expand_document, coerce_marks,
    md_tokens, workspace_height, label_text, label_sep, table_row_height, preamble_blocks,
    coerce_headers, coerce_rows,
)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_BREAK, WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:  # pragma: no cover
    print("error: python-docx is required — pip install \"python-docx==1.1.2\"",
          file=sys.stderr)
    sys.exit(1)


# ---- styles ------------------------------------------------------------------

def _hex_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    if len(h) in (3, 4):  # expand short hex (#1a7 -> #11aa77); HEX in lesson_common accepts 3-8 chars
        h = "".join(c * 2 for c in h[:3])
    h = (h + "000000")[:6]
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def setup_styles(doc, theme: Theme):
    body = float(theme.raw.get("body_size", 10.5))
    s = doc.styles
    s["Normal"].font.name = "Helvetica"
    s["Normal"].font.size = Pt(body)
    # Pin Normal to the html stylesheet's 8pt-after; otherwise Word's template
    # default of 10pt-after leaks into every plain paragraph.
    s["Normal"].paragraph_format.space_before = Pt(0)
    s["Normal"].paragraph_format.space_after = Pt(8)
    for name, size, bold, color, before, after in [
        ("LC Eyebrow", 9, True, theme.safe("primary"), 0, 2),
        ("LC Title", float(theme.raw.get("title_size", 22)), True, theme.safe("text"), 2, 2),
        ("LC Meta", 9, False, theme.safe("muted"), 0, 12),
        ("LC H1", 14, True, theme.safe("text"), 18, 8),
        ("LC H2", 12, True, theme.safe("primary"), 12, 4),
        ("LC H3", body, True, theme.safe("text"), 8, 4),
        ("LC Instr", body, False, theme.safe("text"), 4, 6),
        ("LC Muted", 8.5, False, theme.safe("muted"), 4, 0),
        ("LC Spacer", 2, False, theme.safe("muted"), 4, 0),
    ]:
        st = s.add_style(name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        st.base_style = s["Normal"]
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = _hex_rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
    s["LC Eyebrow"].font.all_caps = True
    # The structural paragraph dropped after every table-rendered block (callout,
    # cards, workspace, table) — OOXML wants a paragraph between consecutive
    # tables. Exact 2pt line height keeps it from reading as a blank line.
    s["LC Spacer"].paragraph_format.line_spacing = Pt(2)


# ---- inline runs -------------------------------------------------------------

def add_md(para, text, bold=False, italic=False):
    """Add mini-markdown text to a paragraph as runs."""
    for t, attrs in md_tokens(text):
        if attrs.get("break"):
            para.add_run().add_break()
            continue
        r = para.add_run(t)
        r.bold = bold or attrs.get("bold", False)
        r.italic = italic or attrs.get("italic", False)


# ---- table primitives --------------------------------------------------------

def _set_borders(tbl, color="CBD2D8", sides=("top", "bottom", "left", "right",
                                             "insideH", "insideV")):
    pr = tbl._tbl.tblPr
    borders = pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        pr.append(borders)
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), color)
        borders.append(el)


def _set_cell_margins(tbl):
    """Internal cell padding (twips; 100 twips = 5pt). Word's default top/bottom is 0,
    which makes cell text sit flush against the top border in most viewers."""
    pr = tbl._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", 100), ("bottom", 100), ("left", 120), ("right", 120)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    pr.append(mar)


def _cant_split(tbl):
    """Mark every row as non-splitting so a table (callout, card row, data row) never
    breaks across a page boundary mid-row."""
    for row in tbl.rows:
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))


def _shade_cell(cell, fill="F6F8F9"):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _list_number_abstract_id(doc) -> str | None:
    """Return the abstractNumId that the built-in 'List Number' style binds to."""
    try:
        # _Cell has no .styles; resolve via the part (cell.part is the DocumentPart).
        styles = doc.styles if hasattr(doc, "styles") else doc.part.document.styles
        st = styles["List Number"].element
        nid = st.find(qn("w:pPr")).find(qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
        for n in doc.part.numbering_part.element.findall(qn("w:num")):
            if n.get(qn("w:numId")) == nid:
                return n.find(qn("w:abstractNumId")).get(qn("w:val"))
    except (AttributeError, KeyError):
        pass
    return None


def _restart_numbering(doc, paras: list):
    """Give an ordered list its own numId so it restarts at 1 instead of continuing
    the previous one (python-docx's built-in 'List Number' style shares one numId
    document-wide). Stamps EVERY paragraph in the list."""
    abstract_id = _list_number_abstract_id(doc)
    if abstract_id is None:
        return
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    new_id = max(int(n.get(qn("w:numId"))) for n in nums) + 1
    new = OxmlElement("w:num")
    new.set(qn("w:numId"), str(new_id))
    abst = OxmlElement("w:abstractNumId")
    abst.set(qn("w:val"), abstract_id)
    new.append(abst)
    ovr = OxmlElement("w:lvlOverride")
    ovr.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    ovr.append(start)
    new.append(ovr)
    numbering.append(new)
    for p in paras:
        npr = p._p.get_or_add_pPr().get_or_add_numPr()
        npr.get_or_add_ilvl().val = 0
        npr.get_or_add_numId().val = new_id


def _keep_with_next(para):
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.keep_together = True


def _box(doc, theme: Theme):
    """A 1x1 bordered table used for callouts."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    _set_borders(t, color=theme.safe("border").lstrip("#"),
                 sides=("top", "bottom", "left", "right"))
    _set_cell_margins(t)
    _cant_split(t)
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    return cell


def _content_width(doc) -> float:
    """Usable text width in inches (page width minus margins)."""
    s = doc.sections[0]
    return (s.page_width - s.left_margin - s.right_margin) / 914400  # EMU per inch


# ---- per-block emitters ------------------------------------------------------

def _emit_paragraph(doc, blk, theme):
    add_md(doc.add_paragraph(), blk.get("text", ""))


def _emit_labeled(doc, blk, theme):
    p = doc.add_paragraph()
    lbl = label_text(blk)
    add_md(p, f"{lbl}{label_sep(lbl)} ", bold=True)
    add_md(p, blk.get("text", ""))


def _emit_heading(style):
    def _f(doc, blk, theme):
        doc.add_paragraph(style=style).add_run(str(blk.get("text", "")))
    return _f


def _emit_instructions(doc, blk, theme):
    add_md(doc.add_paragraph(style="LC Instr"), blk.get("text", ""))


def _label_para(doc, blk):
    # Hug the block it introduces (html .listlabel: 2pt below).
    lp = doc.add_paragraph()
    add_md(lp, label_text(blk), bold=True)
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after = Pt(2)


def _emit_list(doc, blk, theme, *, checklist=False):
    if blk.get("label"):
        _label_para(doc, blk)
    ordered = bool(blk.get("ordered"))
    style = "List Number" if ordered else "List Bullet"
    paras = []
    for item in blk.get("items", []):
        p = doc.add_paragraph(style=style)
        if checklist:
            p.add_run("☐  ")
        add_md(p, item)
        paras.append(p)
    if ordered and paras:
        _restart_numbering(doc, paras)


def _emit_callout(doc, blk, theme):
    kind = resolve_callout_kind(blk)
    icon, _ = CALLOUT_KINDS[kind]
    cell = _box(doc, theme)
    lp = cell.paragraphs[0]
    lp.add_run(icon + "  ")
    if blk.get("label"):
        add_md(lp, label_text(blk), bold=True)
    text = blk.get("text") or blk.get("body") or blk.get("content") or ""
    if text:
        add_md(cell.add_paragraph(), text)
    doc.add_paragraph(style="LC Spacer")


def _emit_cards(doc, blk, theme):
    items = blk.get("items", [])
    if not items:
        return
    tbl = doc.add_table(rows=1, cols=len(items))
    _set_borders(tbl, color=theme.safe("border").lstrip("#"))
    _set_cell_margins(tbl)
    _cant_split(tbl)
    for i, c in enumerate(items):
        c = c if isinstance(c, dict) else {"title": str(c), "text": ""}
        cell = tbl.rows[0].cells[i]
        add_md(cell.paragraphs[0], c.get("title", ""), bold=True)
        add_md(cell.add_paragraph(), c.get("text", ""))
    doc.add_paragraph(style="LC Spacer")


def _emit_fill_in(doc, blk, theme):
    n = FILL_IN_CHARS.get(str(blk.get("size", "med")).lower(), FILL_IN_CHARS["med"])
    p = doc.add_paragraph()
    if blk.get("label"):
        add_md(p, label_text(blk) + ": ", bold=True)
    p.add_run("_" * n)


def _emit_phase_header(doc, blk, theme):
    p = doc.add_paragraph(style="LC H2")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(theme.content_width),
                                              WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(str(blk.get("name", "")))
    if blk.get("minutes") is not None:
        r = p.add_run(f"\t{blk['minutes']} min")
        r.font.bold = False
        r.font.color.rgb = _hex_rgb(theme.safe("muted"))


def _emit_group(doc, blk, theme):
    # A prompt or label never strands from the surface below it: text-bearing
    # paragraphs keep with what follows. Spacers do NOT chain — a multi-surface
    # question (two number lines + a box, 400pt+) as one atomic unit forces Word
    # to jump whole questions to fresh pages, stranding half-blank pages.
    before = len(doc.paragraphs) if hasattr(doc, "paragraphs") else None
    for b in blk.get("blocks", []):
        emit_block(doc, b, theme)
    if before is not None:
        for p in doc.paragraphs[before:-1]:
            if p.text.strip():
                _keep_with_next(p)


def _emit_workspace(doc, blk, theme, *, labeled=False):
    if labeled and blk.get("label"):
        _label_para(doc, blk)
    h = workspace_height(blk, theme)
    if blk.get("ruled", theme.ruled_default):
        # Ruled handwriting lines: one bottom-bordered table row per line — the
        # same border mechanism the number lines use. Paragraph borders (pBdr)
        # never rendered reliably in Word: property-order sensitive, and
        # adjacent identically-bordered paragraphs merge into one box.
        gap = float(theme.answer_gap)
        n = max(2, int(round(h / gap)))
        tbl = doc.add_table(rows=n, cols=1)
        tbl.autofit = False
        _cant_split(tbl)
        for row in tbl.rows:
            row.height = Pt(gap)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            _cell_borders(row.cells[0], ["bottom"], color="C9CFD4", sz="6")
    else:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.rows[0].height = Pt(h)
        tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _set_borders(tbl, color="FFFFFF")
    doc.add_paragraph(style="LC Spacer")


def _emit_page_break(doc, blk, theme):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _emit_table(doc, blk, theme):
    headers = coerce_headers(blk.get("headers"))
    rows = coerce_rows(blk.get("rows"))
    ncols = max(len(headers), max((len(r) for r in rows), default=1), 1)
    tbl = doc.add_table(rows=0, cols=ncols)
    _set_borders(tbl, color=theme.safe("border").lstrip("#"))
    _set_cell_margins(tbl)
    if headers:
        hr = tbl.add_row()
        hr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for i, h in enumerate(headers):
            _shade_cell(hr.cells[i])
            add_md(hr.cells[i].paragraphs[0], str(h).rstrip(": "), bold=True)
    large = blk.get("display") == "large"
    for r in rows:
        tr = tbl.add_row()
        cells = [str(r[i]) if i < len(r) else "" for i in range(ncols)]
        # Underscore runs are write-in blanks, not text.
        cells = ["" if c.strip().strip("_") == "" and "_" in c else c for c in cells]
        full_blank = not any(c.strip() for c in cells)
        for i, v in enumerate(cells):
            p = tr.cells[i].paragraphs[0]
            add_md(p, v, bold=(large and v.strip() != "")
                   or (not headers and i == 0 and v.strip() != ""))
            if large:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(18)
        if any(not c.strip() for c in cells):
            tr.height = Pt(table_row_height(blk, theme, full_blank=full_blank))
            tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _cant_split(tbl)
    doc.add_paragraph(style="LC Spacer")


def _emit_columns(doc, blk, theme):
    tbl = doc.add_table(rows=1, cols=2)
    _set_borders(tbl, color="FFFFFF")
    # A side-by-side comparison is one visual unit — never let Word split the row
    # across a page boundary (a number line half on each page is unreadable).
    _cant_split(tbl)
    for side, cell in (("left", tbl.rows[0].cells[0]), ("right", tbl.rows[0].cells[1])):
        for b in blk.get(side, []):
            emit_block(cell, b, theme)


def _emit_source_card(doc, blk, theme):
    head = " · ".join(str(blk.get(k))
                      for k in ("title", "author", "date", "origin") if blk.get(k))
    _emit_callout(doc, {"kind": "student-task", "label": head,
                        "text": blk.get("excerpt") or blk.get("text") or ""}, theme)


def _emit_fill_table(doc, blk, theme):
    headers = coerce_headers(blk.get("headers"))
    try:
        cols = max(1, len(headers) or int(blk.get("cols") or 2))
    except (TypeError, ValueError):
        cols = 2
    cols = min(cols, 12)
    rows_val = blk.get("rows")
    if isinstance(rows_val, list):
        # Mixed rows: a non-empty list renders its cells (a worked example);
        # an empty list [] renders a blank write-in row.
        rows = []
        for r in coerce_rows(rows_val[:50]):
            cells = r[:cols]
            rows.append(cells + [""] * (cols - len(cells)))
    else:
        try:
            n = int(blk.get("blank_rows") or rows_val or 3)
        except (TypeError, ValueError):
            n = 3
        rows = [[""] * cols for _ in range(min(max(1, n), 50))]
    fwd = {"headers": headers, "rows": rows}
    for k in ("row_height_pt", "empty_row_height_pt"):
        if blk.get(k):
            fwd[k] = blk[k]
    _emit_table(doc, fwd, theme)


def _cell_borders(cell, sides, color="3A4046", sz="12"):
    """Per-cell borders (the table-level helper paints every cell the same)."""
    pr = cell._tc.get_or_add_tcPr()
    borders = pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        pr.append(borders)
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)


def _emit_number_line(doc, blk, theme):
    """A drawn number line, same geometry as the html twin: a horizontal bar with
    evenly spaced tick marks and the min/max labeled at the ends. Built from a
    borderless table — the segment cells' bottom borders form the line, their
    left/right borders form the ticks — so students mark values on it by hand.

    ticks=0 is a distinct "blank line" for student partitioning: the bar and its
    end labels still draw, but with no tick marks at all, not even at the ends."""
    try:
        raw_ticks = None if blk.get("ticks") is None else int(blk.get("ticks"))
    except (TypeError, ValueError):
        raw_ticks = None
    is_blank = raw_ticks == 0
    ticks = 10 if raw_ticks is None else raw_ticks
    ticks = min(max(1, ticks), 24)  # one cell per segment; 24 is plenty on paper
    lo, hi = blk.get("min", 0), blk.get("max", 10)
    marks = coerce_marks(blk.get("marks"))
    try:
        lo_f, hi_f = float(lo), float(hi)
        span = hi_f - lo_f or 1.0
    except (TypeError, ValueError):
        lo_f, hi_f, span = 0.0, float(ticks), float(ticks)
    eps = abs(span) * 0.002
    marks = [(v, lab) for v, lab in marks
             if abs(v - lo_f) > eps and abs(v - hi_f) > eps]

    line_color = theme.safe("text").lstrip("#")
    rows = 3 if marks else 2
    # A single segment (ticks==1, or ticks==0 clamped to one) has only one bar but
    # still needs two independently positioned end labels. A table cell can't hold
    # two independently aligned runs reliably across docx viewers (a tab stop
    # inside a merged cell doesn't render everywhere), so we always give the
    # label row two real grid columns and merge them back into one cell for the
    # bar/marks rows, which every docx viewer renders correctly.
    single_segment = ticks == 1
    cols = 2 if single_segment else ticks
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = False
    _cant_split(tbl)
    # cantSplit stops mid-row breaks; keep-with-next on every paragraph above the
    # last row stops Word separating the line from its labels between rows.
    for row in tbl.rows[:-1]:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True
    r = 0
    if marks:
        # A mark row above the line: label + ▼ centered over the nearest segment boundary.
        mark_row = tbl.rows[0]
        if single_segment:
            mark_row.cells[0].merge(mark_row.cells[1])
        for v, lab in marks:
            seg = 0 if single_segment else min(ticks - 1, max(0, int(round((v - lo_f) / span * ticks))))
            p = mark_row.cells[seg].paragraphs[0]
            if lab:
                r_lab = p.add_run(str(lab) + " ")
                r_lab.bold = True
                r_lab.font.size = Pt(8.5)
            p.add_run("▼").bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
        r = 1
    # The line row: bottom borders draw the bar, left/right borders draw the end
    # ticks — omitted for a blank line, which shows only the bar and its labels.
    line = tbl.rows[r]
    line.height = Pt(22)  # room above the line for student marks
    line.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    if single_segment:
        line_cell = line.cells[0].merge(line.cells[1])
        sides = ["bottom"] + ([] if is_blank else ["left", "right"])
        _cell_borders(line_cell, sides, color=line_color)
        line_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    else:
        for i, cell in enumerate(line.cells):
            sides = ["bottom", "left"] + (["right"] if i == ticks - 1 else [])
            _cell_borders(cell, sides, color=line_color)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    # The label row: min under the first tick, max under the last — always two
    # distinct grid cells, even for a single segment (see `cols` above).
    labels = tbl.rows[r + 1]
    lp = labels.cells[0].paragraphs[0]
    add_md(lp, str(lo))
    lp.paragraph_format.space_after = Pt(0)
    rp = labels.cells[cols - 1].paragraphs[0]
    add_md(rp, str(hi))
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    doc.add_paragraph(style="LC Spacer")


# Adding a block type or text field? Render it in render_lesson_html.py in the
# same commit — the html and docx renderers must emit the same text.
_EMITTERS = {
    "paragraph": _emit_paragraph,
    "labeled": _emit_labeled,
    "h2": _emit_heading("LC H2"),
    "h3": _emit_heading("LC H3"),
    "instructions": _emit_instructions,
    "list": _emit_list,
    "checklist": lambda d, b, t: _emit_list(d, b, t, checklist=True),
    "callout": _emit_callout,
    "cards": _emit_cards,
    "fill_in": _emit_fill_in,
    "phase_header": _emit_phase_header,
    "group": _emit_group,
    "workspace": _emit_workspace,
    "labeled_box": lambda d, b, t: _emit_workspace(d, b, t, labeled=True),
    "page_break": _emit_page_break,
    "table": _emit_table,
    "columns": _emit_columns,
    "source_card": _emit_source_card,
    "fill_table": _emit_fill_table,
    "number_line": _emit_number_line,
}


def emit_block(doc, blk: dict, theme: Theme):
    fn = _EMITTERS.get(_btype(blk))
    if fn is not None:
        return fn(doc, blk, theme)
    if blk.get("text"):
        add_md(doc.add_paragraph(), blk["text"])
    elif blk.get("items"):
        for item in blk["items"]:
            add_md(doc.add_paragraph(style="List Bullet"), item)


# ---- document assembly -------------------------------------------------------

def render(data: dict, out_path: str) -> int:
    data = expand_document(data, data.get("audience", "teacher"))
    theme = Theme(data.get("theme"))
    (theme.answer_height, theme.answer_gap,
     theme.answer_row, theme.ruled_default) = answer_profile(data)
    theme.student_doc = data.get("audience") == "student"

    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = s.bottom_margin = Inches(0.6)
    theme.content_width = _content_width(doc)
    setup_styles(doc, theme)

    hdr = build_header(data)
    if hdr["eyebrow"]:
        doc.add_paragraph(hdr["eyebrow"], style="LC Eyebrow")
    doc.add_paragraph(hdr["title"], style="LC Title")
    if hdr["meta"]:
        doc.add_paragraph(hdr["meta"], style="LC Meta")
    if hdr["name_line"]:
        p = doc.add_paragraph()
        add_md(p, hdr["name_line"])
        p.paragraph_format.space_after = Pt(16)

    for blk in preamble_blocks(data):
        emit_block(doc, blk, theme)

    for section in data.get("sections", []):
        h1 = doc.add_paragraph(style="LC H1")
        h1.paragraph_format.keep_with_next = True
        h1.add_run(str(section.get("heading", "")).rstrip(": "))
        for blk in section.get("blocks", []):
            emit_block(doc, blk, theme)

    if data.get("footer_note"):
        doc.add_paragraph(str(data["footer_note"]), style="LC Muted")

    doc.save(out_path)
    # Post-expansion count, so multi-document sources report accurately.
    return len(data.get("sections", []))


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input", help="lesson JSON file")
    ap.add_argument("-o", "--output", default="lesson_plan.docx")
    args = ap.parse_args()
    data = json.loads(Path(args.input).read_text(encoding="utf-8"))
    render(data, args.output)
    print(f"wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
